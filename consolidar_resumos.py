import os
import json
from docx import Document
from tqdm import tqdm

def consolidar_resumos():
    # Caminhos
    pasta_resumos = 'resumos'
    arquivo_saida = 'resumos_consolidados.docx'
    
    # Verifica se a pasta existe
    if not os.path.exists(pasta_resumos):
        print(f"Erro: A pasta '{pasta_resumos}' não foi encontrada.")
        return

    # Inicializa o documento Word
    doc = Document()
    doc.add_heading('Consolidação de Resumos', 0)

    # Lista arquivos JSON e ordena para manter consistência
    arquivos = sorted([f for f in os.listdir(pasta_resumos) if f.endswith('.json')])
    
    if not arquivos:
        print("Nenhum arquivo JSON encontrado na pasta 'resumos'.")
        return

    print(f"Processando {len(arquivos)} arquivos...")

    for nome_arquivo in tqdm(arquivos, desc="Consolidando"):
        caminho_arquivo = os.path.join(pasta_resumos, nome_arquivo)
        
        try:
            with open(caminho_arquivo, 'r', encoding='utf-8') as f:
                dados = json.load(f)
                
                # Pega o campo 'resposta'
                resposta = dados.get('resposta', 'Campo "resposta" não encontrado.')
                
                # Adiciona ao documento
                # Podemos adicionar uma linha horizontal ou o nome do arquivo para separar,
                # mas o usuário pediu apenas o resumo. Vou adicionar o resumo como um parágrafo.
                p = doc.add_paragraph(resposta)
                
                # Adiciona um espaço entre resumos
                doc.add_paragraph("")
                
        except Exception as e:
            print(f"Erro ao processar {nome_arquivo}: {e}")

    # Salva o arquivo final
    doc.save(arquivo_saida)
    print(f"\nSucesso! Arquivo '{arquivo_saida}' criado com sucesso.")

if __name__ == "__main__":
    consolidar_resumos()
